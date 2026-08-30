import streamlit as st
import sqlite3
import pandas as pd
import os
from datetime import datetime
import pytz
import base64
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageOps

# ---------------------------------------------------------
# 1. CẤU HÌNH TRANG & ICON
# ---------------------------------------------------------
def get_base64_of_bin_file(bin_file):
    try:
        with open(bin_file, 'rb') as f:
            data = f.read()
        return base64.b64encode(data).decode()
    except:
        return ''

icon_base64 = get_base64_of_bin_file('icon.png')
icon_link = f'data:image/png;base64,{icon_base64}' if icon_base64 else ''

st.set_page_config(
    page_title="Hệ Thống Báo Cáo Vận Hành Kho & Thiết Bị", 
    page_icon="icon.png" if os.path.exists("icon.png") else "❄️", 
    layout="wide"
)

if icon_link:
    st.markdown(
        f"""
        <head>
            <link rel="apple-touch-icon" sizes="192x192" href="{icon_link}">
            <link rel="apple-touch-icon-precomposed" href="{icon_link}">
            <link rel="icon" type="image/png" sizes="192x192" href="{icon_link}">
            <link rel="shortcut icon" href="{icon_link}">
        </head>
        """,
        unsafe_allow_html=True
    )

# HÀM NÉN & XỬ LÝ ẢNH
def process_and_save_image(file_buffer, save_path):
    try:
        img = Image.open(file_buffer)
        img = ImageOps.exif_transpose(img)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        img.thumbnail((800, 800))
        img.save(save_path, "JPEG", optimize=True, quality=75)
        return True
    except Exception as e:
        st.error(f"Lỗi xử lý ảnh: {e}")
        return False

# HÀM KẺ KHUNG BẢNG WORD (ĐÃ L BỎ CỘT ĐỘ ẨM)
def format_word_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '555555')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    for i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tcPr = cell._tc.get_or_add_tcPr()
            if i == 0:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '0288D1')
                tcPr.append(shd)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)

# ---------------------------------------------------------
# 2. KHỞI TẠO CƠ SỞ DỮ LIỆU SẠCH
# ---------------------------------------------------------
UPLOAD_DIR = "uploads"
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)

MASTER_DOCX_PATH = os.path.join(UPLOAD_DIR, "Baocao_Tonghop_Capnhat.docx")

conn = sqlite3.connect("kho_system_v104.db", check_same_thread=False)
cursor = conn.cursor()

cursor.execute('''
CREATE TABLE IF NOT EXISTS bao_cao_tong_hop (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    ma_ca_truc TEXT UNIQUE,
    thoi_gian TEXT,
    ca_truc TEXT,
    nguoi_bao_cao TEXT,
    ghi_chu_chung TEXT
)
''')

cursor.execute('''
CREATE TABLE IF NOT EXISTS chi_tiet_kho (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    ma_ca_truc TEXT,
    ten_kho TEXT,
    nhiet_do REAL,
    san_luong REAL,
    trang_thai_may TEXT,
    ghi_chu_kho TEXT,
    duong_dan_anh TEXT
)
''')

# Danh mục Kho
cursor.execute('CREATE TABLE IF NOT EXISTS danh_muc_kho (id INTEGER PRIMARY KEY AUTOINCREMENT, ten_kho TEXT UNIQUE)')
if cursor.execute('SELECT COUNT(*) FROM danh_muc_kho').fetchone()[0] == 0:
    danh_sach_kho_ban_dau = [f"Cụm Kho Số {i}" for i in range(1, 8)] + ["Phòng Máy Nén", "Trạm Biến Áp"]
    for kho in danh_sach_kho_ban_dau:
        cursor.execute('INSERT OR IGNORE INTO danh_muc_kho (ten_kho) VALUES (?)', (kho,))

# Danh mục Trạng thái
cursor.execute('CREATE TABLE IF NOT EXISTS danh_muc_trang_thai (id INTEGER PRIMARY KEY AUTOINCREMENT, ten_trang_thai TEXT UNIQUE)')
if cursor.execute('SELECT COUNT(*) FROM danh_muc_trang_thai').fetchone()[0] == 0:
    ds_tt_ban_dau = ["Bình thường", "Cảnh báo nhẹ", "Sự cố - Cần sửa chữa", "Bảo trì định kỳ"]
    for tt in ds_tt_ban_dau:
        cursor.execute('INSERT OR IGNORE INTO danh_muc_trang_thai (ten_trang_thai) VALUES (?)', (tt,))

# Danh mục Tài khoản
cursor.execute('''
CREATE TABLE IF NOT EXISTS tai_khoan (
    username TEXT PRIMARY KEY,
    password TEXT,
    ho_ten TEXT,
    vai_tro TEXT,
    trang_thai TEXT
)
''')

cursor.execute('INSERT OR IGNORE INTO tai_khoan (username, password, ho_ten, vai_tro, trang_thai) VALUES ("admin", "111111", "Quản Trị Viên Nguyễn Trọng Hưng", "admin", "hoat_dong")')
conn.commit()

# ---------------------------------------------------------
# 3. HÀM TẢI FILE VÀ XUẤT FILE WORD
# ---------------------------------------------------------
def render_download_button(file_path, button_text, filename):
    if os.path.exists(file_path):
        with open(file_path, "rb") as f:
            bytes_data = f.read()
        b64 = base64.b64encode(bytes_data).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" target="_blank" style="text-decoration:none;"><button style="background-color:#0288d1;color:white;padding:10px 15px;border:none;border-radius:8px;font-weight:bold;cursor:pointer;width:100%;">📥 {button_text}</button></a>'
        st.markdown(href, unsafe_allow_html=True)

def add_multi_image_grid_to_docx(doc, chi_tiet_items):
    doc.add_paragraph("")
    h = doc.add_heading("📸 Hình Ảnh Báo Cáo Chi Tiết Theo Từng Kho:", level=2)
    h.paragraph_format.space_after = Pt(6)

    for item in chi_tiet_items:
        ten_kho = item[0]
        nhiet_do = item[1]
        trang_thai = item[3]
        raw_paths = item[5]
        
        if raw_paths:
            img_paths = [p.strip() for p in raw_paths.split(";") if p.strip() and os.path.exists(p.strip())]
            if img_paths:
                p_head = doc.add_paragraph()
                r_head = p_head.add_run(f"📌 {ten_kho} (Nhiệt độ: {nhiet_do}°C | TT: {trang_thai}) - {len(img_paths)} hình ảnh:")
                r_head.bold = True
                
                grid_table = doc.add_table(rows=0, cols=2)
                grid_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                for idx_img in range(0, len(img_paths), 2):
                    row_cells = grid_table.add_row().cells
                    
                    p_l = row_cells[0].paragraphs[0]
                    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        p_l.add_run().add_picture(img_paths[idx_img], width=Inches(2.6))
                    except:
                        p_l.add_run("[Lỗi ảnh]")
                        
                    if idx_img + 1 < len(img_paths):
                        p_r = row_cells[1].paragraphs[0]
                        p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        try:
                            p_r.add_run().add_picture(img_paths[idx_img+1], width=Inches(2.6))
                        except:
                            p_r.add_run("[Lỗi ảnh]")

def append_to_master_docx(ma_ca):
    row_ca = cursor.execute("SELECT * FROM bao_cao_tong_hop WHERE ma_ca_truc = ?", (ma_ca,)).fetchone()
    chi_tiet = cursor.execute("SELECT ten_kho, nhiet_do, san_luong, trang_thai_may, ghi_chu_kho, duong_dan_anh FROM chi_tiet_kho WHERE ma_ca_truc = ?", (ma_ca,)).fetchall()

    if os.path.exists(MASTER_DOCX_PATH):
        doc = Document(MASTER_DOCX_PATH)
        doc.add_page_break()
    else:
        doc = Document()
        doc.add_heading("SỔ TAY TỔNG HỢP BÁO CÁO VẬN HÀNH KHO & THIẾT BỊ", level=0)

    doc.add_heading(f"📌 BÁO CÁO CA TRỰC: {ma_ca}", level=1)
    doc.add_paragraph(f"⏰ Thời gian gửi báo cáo: {row_ca[2]} | 🔄 Ca / Khung giờ: {row_ca[3]} | 👤 Người báo cáo: {row_ca[4]}")
    doc.add_paragraph(f"📝 Ghi chú chung: {row_ca[5] if row_ca[5] else 'Không có'}")
    
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = 'Kho / Thiết Bị', 'Nhiệt Độ (°C)', 'Sản Lượng (Tấn)', 'Trạng Thái', 'Ghi Chú Riêng'
    
    for item in chi_tiet:
        r = table.add_row().cells
        r[0].text, r[1].text, r[2].text, r[3].text, r[4].text = str(item[0]), str(item[1]), str(item[2]), str(item[3]), str(item[4] if item[4] else "")

    format_word_table(table)
    add_multi_image_grid_to_docx(doc, chi_tiet)
    doc.save(MASTER_DOCX_PATH)

    # FILE RIÊNG CA TRỰC
    single_doc_path = os.path.join(UPLOAD_DIR, f"BaoCao_{ma_ca}.docx")
    single_doc = Document()
    single_doc.add_heading(f"BÁO CÁO CA TRỰC: {ma_ca}", level=1)
    single_doc.add_paragraph(f"⏰ Thời gian gửi: {row_ca[2]} | Ca / Khung giờ: {row_ca[3]} | Người báo cáo: {row_ca[4]}")
    single_doc.add_paragraph(f"Ghi chú chung: {row_ca[5] if row_ca[5] else 'Không có'}")
    
    t_s = single_doc.add_table(rows=1, cols=5)
    h_s = t_s.rows[0].cells
    h_s[0].text, h_s[1].text, h_s[2].text, h_s[3].text, h_s[4].text = 'Kho / Thiết Bị', 'Nhiệt độ (°C)', 'Sản lượng (Tấn)', 'Trạng thái', 'Ghi chú riêng'
    for item in chi_tiet:
        rs = t_s.add_row().cells
        rs[0].text, rs[1].text, rs[2].text, rs[3].text, rs[4].text = str(item[0]), str(item[1]), str(item[2]), str(item[3]), str(item[4] if item[4] else "")
    
    format_word_table(t_s)
    add_multi_image_grid_to_docx(single_doc, chi_tiet)
    single_doc.save(single_doc_path)

# ---------------------------------------------------------
# 4. ĐĂNG NHẬP & PHÂN QUYỀN
# ---------------------------------------------------------
if "user_info" not in st.session_state:
    st.session_state["user_info"] = None

if "admin_unlocked" not in st.session_state:
    st.session_state["admin_unlocked"] = False

st.sidebar.image("https://img.icons8.com/color/96/000000/cold-storage.png", width=70)
st.sidebar.title("🔐 ĐĂNG NHẬP")

with st.sidebar:
    if st.session_state["user_info"] is None:
        u_input = st.text_input("Tài khoản").strip()
        p_input = st.text_input("Mật khẩu (NV để trống)", type="password").strip()
        if st.button("Đăng nhập"):
            user_query = cursor.execute("SELECT username, password, ho_ten, vai_tro, trang_thai FROM tai_khoan WHERE username = ?", (u_input,)).fetchone()
            if user_query:
                if user_query[4] == "bi_khoa":
                    st.error("❌ Tài khoản này đã bị khóa!")
                elif user_query[1] == p_input:
                    st.session_state["user_info"] = {"username": user_query[0], "ho_ten": user_query[2], "vai_tro": user_query[3]}
                    if user_query[3] == "admin":
                        st.session_state["admin_unlocked"] = True
                    st.success(f"Xin chào: {user_query[2]}")
                    st.rerun()
                else:
                    st.error("Sai mật khẩu!")
            else:
                st.error("Tài khoản không tồn tại!")
    else:
        u_info = st.session_state["user_info"]
        st.success(f"👤 **{u_info['ho_ten']}**")
        st.caption(f"Quyền: **{u_info['vai_tro'].upper()}**")
        if st.button("Đăng xuất"):
            st.session_state["user_info"] = None
            st.session_state["admin_unlocked"] = False
            st.rerun()

# ---------------------------------------------------------
# 5. GIAO DIỆN CHÍNH
# ---------------------------------------------------------
st.title("❄️ QUẢN LÝ & BÁO CÁO VẬN HÀNH KHO / THIẾT BỊ")

if st.session_state["user_info"] is None:
    st.info("👈 Vui lòng đăng nhập tài khoản ở menu bên trái.")
else:
    current_user = st.session_state["user_info"]
    can_report = current_user["vai_tro"] in ["admin", "nhanvien"]

    ds_kho = [row[0] for row in cursor.execute("SELECT ten_kho FROM danh_muc_kho ORDER BY id ASC").fetchall()]
    ds_tt = [row[0] for row in cursor.execute("SELECT ten_trang_thai FROM danh_muc_trang_thai ORDER BY id ASC").fetchall()]

    tabs = st.tabs(["📊 Xem & Tải Báo Cáo", "📝 Lập Báo Cáo Ca Trực", "⚙️ Cài Đặt Sâu (Chỉ Admin)"])

    # TAB 1: XEM & TẢI BÁO CÁO
    with tabs[0]:
        st.subheader("📊 Nhật Ký Báo Cáo Ca Trực Tổng Hợp")
        
        if os.path.exists(MASTER_DOCX_PATH):
            render_download_button(
                MASTER_DOCX_PATH, 
                "TẢI FILE TỔNG TẤT CẢ BÁO CÁO", 
                "Baocao_Tonghop_Capnhat.docx"
            )
        
        st.markdown("---")
        df_ca = pd.read_sql_query("SELECT * FROM bao_cao_tong_hop ORDER BY id DESC", conn)

        if not df_ca.empty:
            for idx, row_ca in df_ca.iterrows():
                ma_ca = row_ca['ma_ca_truc']
                with st.expander(f"📌 Mã Ca: {ma_ca} | Ngày gửi: {row_ca['thoi_gian']} | Ca: {row_ca['ca_truc']} | Người báo cáo: {row_ca['nguoi_bao_cao']}"):
                    st.write(f"**Ghi chú chung:** {row_ca['ghi_chu_chung']}")
                    
                    single_path = os.path.join(UPLOAD_DIR, f"BaoCao_{ma_ca}.docx")
                    if os.path.exists(single_path):
                        render_download_button(
                            single_path,
                            f"Tải riêng file ca {ma_ca}",
                            f"BaoCao_LichSu_{ma_ca}.docx"
                        )
                    
                    if st.button(f"🗑️ Xóa Lịch Sử Ca Này", key=f"del_ca_{ma_ca}"):
                        cursor.execute("DELETE FROM bao_cao_tong_hop WHERE ma_ca_truc = ?", (ma_ca,))
                        cursor.execute("DELETE FROM chi_tiet_kho WHERE ma_ca_truc = ?", (ma_ca,))
                        conn.commit()
                        st.success(f"Đã xóa lịch sử ca {ma_ca}!")
                        st.rerun()

                    st.markdown("---")
                    
                    df_full = cursor.execute("SELECT ten_kho, nhiet_do, san_luong, trang_thai_may, ghi_chu_kho, duong_dan_anh FROM chi_tiet_kho WHERE ma_ca_truc = ?", (ma_ca,)).fetchall()
                    
                    view_mode = st.radio("Chế độ xem hình ảnh:", ["🖼️ Dạng Lưới Tất Cả Kho", "🔎 Xem Slide Từng Ảnh Độc Lập"], key=f"mode_{ma_ca}", horizontal=True)

                    if view_mode == "🖼️ Dạng Lưới Tất Cả Kho":
                        cols = st.columns(3)
                        for idx_item, item in enumerate(df_full):
                            with cols[idx_item % 3]:
                                with st.container(border=True):
                                    st.markdown(f"#### 🏭 {item[0]}")
                                    st.write(f"• Nhiệt độ: **{item[1]} °C**")
                                    st.write(f"• Sản lượng: **{item[2]} Tấn**")
                                    st.write(f"• Trạng thái: **{item[3]}**")
                                    if item[4]:
                                        st.write(f"• Ghi chú: *{item[4]}*")
                                    
                                    raw_p = item[5]
                                    if raw_p:
                                        paths = [p.strip() for p in raw_p.split(";") if p.strip() and os.path.exists(p.strip())]
                                        if paths:
                                            st.markdown(f"**Ảnh đính kèm ({len(paths)} ảnh):**")
                                            for p in paths:
                                                st.image(p, use_container_width=True)
                                        else:
                                            st.info("Không có ảnh")
                                    else:
                                        st.info("Không có ảnh")
                    else:
                        all_img_list = []
                        for item in df_full:
                            if item[5]:
                                paths = [p.strip() for p in item[5].split(";") if p.strip() and os.path.exists(p.strip())]
                                for p in paths:
                                    all_img_list.append((p, item))

                        if not all_img_list:
                            st.warning("Ca trực này chưa upload hình ảnh nào.")
                        else:
                            idx_key = f"slide_idx_{ma_ca}"
                            if idx_key not in st.session_state:
                                st.session_state[idx_key] = 0

                            curr_idx = st.session_state[idx_key]
                            curr_img_path, curr_item = all_img_list[curr_idx]

                            c_prev, c_info, c_next = st.columns([1, 4, 1])
                            with c_prev:
                                if st.button("◀ Ảnh Trước", key=f"p_{ma_ca}"):
                                    st.session_state[idx_key] = (curr_idx - 1) % len(all_img_list)
                                    st.rerun()
                            with c_next:
                                if st.button("Ảnh Tiếp ▶", key=f"n_{ma_ca}"):
                                    st.session_state[idx_key] = (curr_idx + 1) % len(all_img_list)
                                    st.rerun()

                            with c_info:
                                st.caption(f"Đang xem hình {curr_idx + 1} / {len(all_img_list)}")

                            col_img, col_txt = st.columns([3, 2])
                            with col_img:
                                st.image(curr_img_path, caption=f"Ảnh tại {curr_item[0]}", use_container_width=True)
                            with col_txt:
                                with st.container(border=True):
                                    st.markdown(f"### 📍 Thông Tin {curr_item[0]}")
                                    st.markdown(f"* **Nhiệt độ:** `{curr_item[1]} °C`")
                                    st.markdown(f"* **Sản lượng:** `{curr_item[2]} Tấn`")
                                    st.markdown(f"* **Trạng thái:** `{curr_item[3]}`")
                                    st.markdown(f"* **Ghi chú riêng:** {curr_item[4] if curr_item[4] else 'Không có'}")

        else:
            st.info("Chưa có báo cáo ca trực nào.")

    # TAB 2: LẬP BÁO CÁO CA TRỰC
    with tabs[1]:
        if can_report:
            st.subheader("📝 Lập Báo Cáo Ca Trực Mới")
            
            if "submitted_ca" in st.session_state:
                submitted_ma_ca = st.session_state['submitted_ca']
                st.success(f"🔔 **ĐÃ GỬI BÁO CÁO THÀNH CÔNG! (Mã ca: {submitted_ma_ca})**")
                
                st.markdown("#### 📄 Tải file Word vừa tạo:")
                single_path_new = os.path.join(UPLOAD_DIR, f"BaoCao_{submitted_ma_ca}.docx")
                if os.path.exists(single_path_new):
                    render_download_button(
                        single_path_new,
                        f"TẢI FILE WORD CA NÀY ({submitted_ma_ca})",
                        f"BaoCao_{submitted_ma_ca}.docx"
                    )
                
                st.markdown("---")
                if st.button("🏠 QUAY VỀ LẬP BÁO CÁO MỚI", use_container_width=True):
                    del st.session_state["submitted_ca"]
                    st.rerun()

            else:
                if not ds_kho:
                    st.warning("⚠️ Chưa có kho nào trong hệ thống.")
                else:
                    tz_vn = pytz.timezone('Asia/Ho_Chi_Minh')
                    now_vn = datetime.now(tz_vn)

                    with st.form("form_nhap_ca"):
                        st.markdown("##### 📌 Thông Tin Ca Trực")
                        col_top1, col_top2 = st.columns(2)
                        
                        ma_ca = col_top1.text_input("Mã Ca Trực", value=f"CA-{now_vn.strftime('%Y%m%d-%H%M%S')}")
                        nguoi_lap = col_top2.text_input("Người Báo Cáo", value=current_user["ho_ten"], disabled=True)

                        st.markdown("---")
                        st.markdown("##### ⏰ Giờ Ca Làm (Nhân viên tự nhập)")
                        ca_truc_str = st.text_input("Nhập giờ ca làm / Khung giờ", value="06h00 - 14h00 (Ca sáng)", placeholder="Ví dụ: 06h00 - 14h00 hoặc Ca 1...")
                        ghi_chu_chung = st.text_area("Ghi chú chung ca trực (nếu có)")

                        st.markdown("---")
                        st.write(f"📋 **NHẬP BÁO CÁO CHO {len(ds_kho)} KHO / THIẾT BỊ HIỆN CÓ:**")

                        kho_inputs = {}
                        for kho in ds_kho:
                            with st.container(border=True):
                                st.markdown(f"### ❄️ {kho}")
                                col_l, col_r = st.columns([3, 2])
                                
                                with col_l:
                                    col_a, col_b = st.columns(2)
                                    # NHIỆT ĐỘ MẶC ĐỊNH LÀ 0.0
                                    n_do = col_a.number_input(f"Nhiệt độ (°C)", value=0.0, step=0.1, key=f"nd_{kho}")
                                    s_luong = col_b.number_input(f"Sản lượng (Tấn)", value=0.0, step=0.1, key=f"sl_{kho}")
                                    t_thai = st.selectbox(f"Trạng thái vận hành", ds_tt if ds_tt else ["Bình thường"], key=f"tt_{kho}")
                                    gc_kho = st.text_input(f"📝 Ghi chú riêng cho [{kho}]", value="", key=f"gc_{kho}")

                                with col_r:
                                    files_img = st.file_uploader(f"📸 Chụp / Tải NHIỀU ÁNH cho [{kho}]", type=["jpg", "png", "jpeg"], accept_multiple_files=True, key=f"img_{kho}")

                                kho_inputs[kho] = {
                                    "nhiet_do": n_do, "san_luong": s_luong,
                                    "trang_thai": t_thai, "ghi_chu_kho": gc_kho, "files_img": files_img
                                }

                        btn_submit = st.form_submit_button("🚀 GỬI TOÀN BỘ BÁO CÁO CA TRỰC")

                    if btn_submit:
                        with st.spinner("⏳ Đang xử lý báo cáo và xuất file Word..."):
                            now_real = datetime.now(pytz.timezone('Asia/Ho_Chi_Minh')).strftime("%d/%m/%Y %H:%M:%S")
                            
                            cursor.execute("INSERT INTO bao_cao_tong_hop (ma_ca_truc, thoi_gian, ca_truc, nguoi_bao_cao, ghi_chu_chung) VALUES (?, ?, ?, ?, ?)",
                                           (ma_ca, now_real, ca_truc_str, nguoi_lap, ghi_chu_chung))
                            
                            for kho, data in kho_inputs.items():
                                saved_img_paths = []
                                if data["files_img"]:
                                    for idx_f, f_img in enumerate(data["files_img"]):
                                        img_filename = f"{ma_ca}_{kho}_{idx_f+1}.jpg"
                                        img_path = os.path.join(UPLOAD_DIR, img_filename)
                                        if process_and_save_image(f_img, img_path):
                                            saved_img_paths.append(img_path)

                                str_img_paths = ";".join(saved_img_paths)

                                cursor.execute('''
                                    INSERT INTO chi_tiet_kho (ma_ca_truc, ten_kho, nhiet_do, san_luong, trang_thai_may, ghi_chu_kho, duong_dan_anh)
                                    VALUES (?, ?, ?, ?, ?, ?, ?)
                                ''', (ma_ca, kho, data["nhiet_do"], data["san_luong"], data["trang_thai"], data["ghi_chu_kho"], str_img_paths))

                            conn.commit()
                            append_to_master_docx(ma_ca)
                            st.session_state["submitted_ca"] = ma_ca
                            st.rerun()
        else:
            st.warning("🔒 Tài khoản của bạn là quyền VIEWER (Chỉ xem), không được phép lập báo cáo.")

    # TAB 3: CÀI ĐẶT SÂU
    with tabs[2]:
        st.subheader("⚙️ Cài Đặt Hệ Thống & Quản Lý Sâu")
        
        if not st.session_state["admin_unlocked"]:
            st.warning("🔒 Khu vực này yêu cầu Mật Khẩu Admin để truy cập cài đặt sâu.")
            admin_pwd_input = st.text_input("Nhập Mật Khẩu Admin để mở khóa:", type="password", key="unlock_admin_pwd")
            if st.button("Mở Khóa Cài Đặt"):
                admin_pass_db = cursor.execute("SELECT password FROM tai_khoan WHERE username = 'admin'").fetchone()
                if admin_pass_db and admin_pwd_input == admin_pass_db[0]:
                    st.session_state["admin_unlocked"] = True
                    st.success("🔓 Xác thực thành công!")
                    st.rerun()
                else:
                    st.error("❌ Mật khẩu Admin không chính xác!")
        else:
            st.success("🔓 Đã mở khóa quyền Admin")
            if st.button("🔒 Khóa Lại Cài Đặt"):
                st.session_state["admin_unlocked"] = False
                st.rerun()

            st.markdown("---")
            sub_tab1, sub_tab2 = st.tabs(["🏬 Quản Lý Kho & Trạng Thái", "👥 Quản Lý Tài Khoản"])

            with sub_tab1:
                st.markdown("### 🏬 1. Tùy Chỉnh Kho / Thiết Bị")
                c_k1, c_k2, c_k3 = st.columns(3)
                
                with c_k1:
                    with st.container(border=True):
                        st.markdown("##### ➕ Thêm Kho Mới")
                        ten_kho_moi = st.text_input("Nhập tên kho mới", key="add_kho_input")
                        if st.button("Thêm Vào Danh Mục", key="btn_add_kho"):
                            if ten_kho_moi.strip():
                                try:
                                    cursor.execute("INSERT INTO danh_muc_kho (ten_kho) VALUES (?)", (ten_kho_moi.strip(),))
                                    conn.commit()
                                    st.success(f"Đã thêm kho: {ten_kho_moi}")
                                    st.rerun()
                                except:
                                    st.error("Tên kho này đã tồn tại!")

                with c_k2:
                    with st.container(border=True):
                        st.markdown("##### ✏️ Đổi Tên Kho")
                        kho_doi = st.selectbox("Chọn kho cần đổi tên", ds_kho if ds_kho else ["Chưa có"], key="sel_doi_kho")
                        ten_kho_renamed = st.text_input("Nhập tên kho mới", key="txt_ren_kho")
                        if st.button("Cập Nhật Tên Kho", key="btn_ren_kho"):
                            if ds_kho and ten_kho_renamed.strip() and ten_kho_renamed.strip() != kho_doi:
                                try:
                                    cursor.execute("UPDATE danh_muc_kho SET ten_kho = ? WHERE ten_kho = ?", (ten_kho_renamed.strip(), kho_doi))
                                    conn.commit()
                                    st.success("Đã đổi tên kho thành công!")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Lỗi đổi tên kho: {e}")

                with c_k3:
                    with st.container(border=True):
                        st.markdown("##### 🗑️ XÓA HẲN KHO")
                        kho_xoa = st.selectbox("Chọn kho cần XÓA MẤT LUÔN", ds_kho if ds_kho else ["Chưa có"], key="sel_xoa_kho")
                        if st.button("❌ XÓA VĨNH VIỄN KHO NÀY", key="btn_del_kho"):
                            if ds_kho:
                                cursor.execute("DELETE FROM danh_muc_kho WHERE ten_kho = ?", (kho_xoa,))
                                conn.commit()
                                st.success(f"Đã xóa vĩnh viễn kho '{kho_xoa}'!")
                                st.rerun()

                st.markdown("---")
                st.markdown("### 🔴 2. Tùy Chỉnh Danh Mục Trạng Thái Vận Hành")
                c_t1, c_t2 = st.columns(2)
                
                with c_t1:
                    with st.container(border=True):
                        st.markdown("##### ➕ Thêm Trạng Thái Mới")
                        tt_moi = st.text_input("Ví dụ: Đang sửa chữa máy nén 2", key="add_tt_input")
                        if st.button("Thêm Trạng Thái", key="btn_add_tt"):
                            if tt_moi.strip():
                                try:
                                    cursor.execute("INSERT INTO danh_muc_trang_thai (ten_trang_thai) VALUES (?)", (tt_moi.strip(),))
                                    conn.commit()
                                    st.success("Đã thêm trạng thái mới!")
                                    st.rerun()
                                except:
                                    st.error("Trạng thái này đã có!")

                with c_t2:
                    with st.container(border=True):
                        st.markdown("##### 🗑️ Xóa Trạng Thái")
                        tt_xoa = st.selectbox("Chọn trạng thái cần xóa", ds_tt if ds_tt else ["Chưa có"], key="sel_xoa_tt")
                        if st.button("❌ Xóa Trạng Thái Này", key="btn_del_tt"):
                            if ds_tt:
                                cursor.execute("DELETE FROM danh_muc_trang_thai WHERE ten_trang_thai = ?", (tt_xoa,))
                                conn.commit()
                                st.success(f"Đã xóa trạng thái '{tt_xoa}'!")
                                st.rerun()

            with sub_tab2:
                st.markdown("### 👥 Cấp Tài Khoản & Quản Lý Nhân Viên")
                
                with st.form("form_tao_tk"):
                    st.markdown("##### ➕ Tạo tài khoản mới")
                    c_u, c_p, c_n, c_r = st.columns(4)
                    new_u = c_u.text_input("Tên Đăng Nhập")
                    new_p = c_p.text_input("Mật Khẩu (để trống nếu muốn)")
                    new_n = c_n.text_input("Họ Và Tên")
                    new_r = c_r.selectbox("Phân Quyền", ["nhanvien", "viewer", "admin"])
                    
                    if st.form_submit_button("TẠO TÀI KHOẢN"):
                        if new_u and new_n:
                            try:
                                cursor.execute("INSERT INTO tai_khoan (username, password, ho_ten, vai_tro, trang_thai) VALUES (?, ?, ?, ?, 'hoat_dong')", (new_u, new_p.strip(), new_n, new_r))
                                conn.commit()
                                st.success(f"Đã tạo tài khoản cho {new_n} ({new_r})!")
                                st.rerun()
                            except:
                                st.error("Tên đăng nhập này đã có người sử dụng!")

                st.markdown("---")
                st.subheader("📋 Danh Sách Tài Khoản Trong Hệ Thống")
                df_users = pd.read_sql_query("SELECT username AS 'Tên ĐN', ho_ten AS 'Họ Tên', vai_tro AS 'Quyền', trang_thai AS 'Trạng Thái' FROM tai_khoan", conn)
                st.dataframe(df_users, use_container_width=True)
                
                st.markdown("##### 🛠️ Thao Tác Với Tài Khoản Nhân Viên")
                col_usr, col_b1, col_b2 = st.columns([2, 1.5, 1.5])
                usr_target = col_usr.selectbox("Chọn tài khoản", df_users['Tên ĐN'].tolist())
                
                if usr_target == "admin":
                    st.info("💡 Tài khoản 'admin' mặc định không thể xóa.")
                else:
                    if col_b1.button("🗑️ XÓA VĨNH VIỄN TÀI KHOẢN"):
                        cursor.execute("DELETE FROM tai_khoan WHERE username = ?", (usr_target,))
                        conn.commit()
                        st.success(f"Đã XÓA MẤT LUÔN tài khoản {usr_target} khỏi hệ thống!")
                        st.rerun()

                    if col_b2.button("🔒 Khóa / Mở Khóa"):
                        curr_st = cursor.execute("SELECT trang_thai FROM tai_khoan WHERE username = ?", (usr_target,)).fetchone()[0]
                        new_st = "bi_khoa" if curr_st == "hoat_dong" else "hoat_dong"
                        cursor.execute("UPDATE tai_khoan SET trang_thai = ? WHERE username = ?", (new_st, usr_target))
                        conn.commit()
                        st.success(f"Đã đổi trạng thái tài khoản {usr_target} thành: {new_st}")
                        st.rerun()